#!/usr/bin/env python3
"""Word audit renderer (.docx) behind python-docx.

Turns the report view into a written audit: an overview, a winnable-clicks section,
a citation-readiness section, a clusters table, an entity-gap table, the observed
section when Live Wire is present, and a how-to-read section. Authorship and the
core-properties dates are pinned to the run timestamp so the file is byte-stable
across runs once the port normalises the archive.

This path needs ``python-docx`` (an optional backend). The caller checks
``ooxml_capabilities()`` and skips this format cleanly when the package is absent.
The rendering is exercised in continuous integration; locally the suite degrades to
the HTML dashboard.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from .brand import office_hex
from .dashboard_html import _fmt_band, _fmt_int, _fmt_pct


def _table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass  # Style set depends on the template; plain grid is acceptable.
    head_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        head_cells[i].text = text
        for para in head_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def render_audit(
    view: dict[str, Any], brand: dict[str, Any], out_path: Path, timestamp: datetime
) -> None:
    """Render the Word audit to ``out_path``. Requires python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    header = view["header"]
    title_text = header["label"] or f"{brand['name']} keyword and demand audit"

    heading = doc.add_heading(title_text, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(
            office_hex(brand["colors"]["primary"])
        )
    sub = doc.add_paragraph(f"{brand['name']} | {brand['tagline']}")
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(office_hex(brand["colors"]["muted"]))

    # Demand overview.
    corpus = view["corpus"]
    doc.add_heading("Demand overview", level=1)
    doc.add_paragraph(f"Keywords analysed: {_fmt_int(corpus['total_keywords'])}")
    doc.add_paragraph(
        f"AI Overview eligible: {_fmt_pct(corpus['aio_eligibility_share'])}"
    )
    doc.add_paragraph(
        f"Generative-engine opportunity: {_fmt_pct(corpus['geo_opportunity_share'])}"
    )
    doc.add_paragraph(f"Demand opportunity score: {corpus['demand_opportunity_score']}")
    _table(
        doc,
        ["Search intent", "Keywords"],
        [
            [r["intent"].replace("_", " "), str(r["count"])]
            for r in corpus["intent_split"]
        ],
    )

    # Winnable clicks.
    winnable = view.get("winnable")
    if winnable:
        doc.add_heading("Winnable clicks", level=1)
        doc.add_paragraph(
            f"Portfolio winnable band per month: {_fmt_band(winnable.get('portfolio_winnable_band'))} "
            f"(modelled current {_fmt_int(winnable.get('current_clicks_estimate'))})."
        )
        if "observed_current_clicks" in winnable:
            doc.add_paragraph(
                f"Observed current clicks (Live Wire): {_fmt_int(winnable.get('observed_current_clicks'))}; "
                f"observed winnable band {_fmt_band(winnable.get('observed_winnable_band'))}."
            )
        _table(
            doc,
            ["Intent", "Winnable clicks (band)"],
            [
                [r["intent"].replace("_", " "), _fmt_band(r["winnable_band"])]
                for r in winnable["by_intent"]
            ],
        )

    # Citation readiness.
    readiness = view.get("readiness")
    if readiness:
        doc.add_heading("AI-citation readiness (expected)", level=1)
        doc.add_paragraph(
            f"Mean expected readiness: {readiness['mean_expected_readiness']} out of 100, across "
            f"{readiness['scored_components']} scored components "
            f"({readiness['checklist_only_components']} checklist-only)."
        )

    # Clusters by demand.
    doc.add_heading("Clusters by demand", level=1)
    cluster_rows = [
        [
            str(c["head"]),
            _fmt_int(c.get("volume_total")),
            str(c.get("dominant_intent", "")).replace("_", " "),
            str(c.get("topical_authority", "n/a")),
            str(c.get("expected_readiness", "n/a")),
            _fmt_band(c.get("winnable_band")),
        ]
        for c in view["top_clusters"]
    ]
    _table(
        doc,
        ["Cluster", "Volume", "Intent", "Authority", "Readiness", "Winnable clicks"],
        cluster_rows,
    )

    # Entity gaps.
    gaps = view.get("gaps") or []
    if gaps:
        doc.add_heading("Entity gaps", level=1)
        _table(
            doc,
            ["Entity", "Demand", "Suggested cluster"],
            [
                [
                    g["entity"],
                    _fmt_int(g.get("demand_volume")),
                    str(g.get("suggested_cluster_head")),
                ]
                for g in gaps
            ],
        )

    # Observed citations.
    obs = view.get("observed")
    if obs:
        doc.add_heading("Observed citations (Live Wire)", level=1)
        surfaces = ", ".join(obs.get("surfaces") or []) or "n/a"
        doc.add_paragraph(
            f"Observed client citation share: {_fmt_pct(obs.get('observed_citation_share'), scale=1.0)} "
            f"on surfaces: {surfaces}."
        )
        _table(
            doc,
            ["Domain", "Citation share"],
            [
                [d["domain"], _fmt_pct(d["share"], scale=1.0)]
                for d in obs.get("competitor_split", [])
            ],
        )

    # How to read.
    notes = view.get("provenance_notes") or []
    if notes:
        doc.add_heading("How to read these figures", level=1)
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")

    footer = doc.add_paragraph(
        f"{brand['footer']} Generated {header['generated_at']}; input hash {header['input_hash'][:12]}."
    )
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(office_hex(brand["colors"]["muted"]))

    core = doc.core_properties
    core.author = brand["author"]
    core.title = title_text
    core.created = timestamp
    core.modified = timestamp

    doc.save(str(out_path))
